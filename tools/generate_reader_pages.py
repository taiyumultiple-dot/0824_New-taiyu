from html import escape
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "assets" / "documents"
OUTPUT = ROOT / "read"
FILES = {
    "閩南語文(上)授課計畫表-1150810.docx": "閩南語文-上-授課計畫表.html",
    "閩南語文(下)授課計畫表-1150810.docx": "閩南語文-下-授課計畫表.html",
    "閩南語文(全)授課計畫表-1150810.docx": "閩南語文-全-授課計畫表.html",
}

STYLE = """body{margin:0;background:#f5f0df;color:#171717;font-family:'Noto Sans TC','Microsoft JhengHei',sans-serif}.bar{padding:18px max(24px,calc((100% - 920px)/2));border-bottom:3px solid #202020;background:#fff}.bar a{color:#171717;font-weight:900;text-decoration:none;border-bottom:2px solid}.page{max-width:920px;margin:40px auto 72px;padding:clamp(24px,5vw,56px);background:#fff;border:3px solid #202020;box-shadow:9px 9px 0 #ffd927}.eyebrow{color:#007f7a;font-size:14px;font-weight:900;letter-spacing:.08em}h1{font-size:clamp(28px,5vw,44px);line-height:1.2;border-bottom:3px solid #202020;padding-bottom:22px}p{line-height:1.85}table{width:100%;border-collapse:collapse;margin:28px 0;font-size:14px}td{padding:10px;border:1px solid #8a8a8a;vertical-align:top;line-height:1.6}@media(max-width:600px){.page{margin:0;border:0;box-shadow:none}.bar{padding:16px 20px}table{font-size:13px}}"""

def render_document(docx_path, title):
    doc = Document(docx_path)
    blocks = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(f"<p>{escape(text)}</p>")
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{escape(cell.text).replace(chr(10), '<br>')}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        blocks.append("<table>" + "".join(rows) + "</table>")
    return f"<!doctype html><html lang='zh-Hant'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width, initial-scale=1'><title>{escape(title)}｜AI 書櫃</title><style>{STYLE}</style></head><body><header class='bar'><a href='../index.html'>← 回到 AI 書櫃</a></header><main class='page'><p class='eyebrow'>第 1 課 AI 書櫃 · 線上閱讀</p><h1>{escape(title)}</h1>{''.join(blocks)}</main></body></html>"

OUTPUT.mkdir(exist_ok=True)
for source_name, output_name in FILES.items():
    title = source_name.replace('-1150810.docx', '')
    (OUTPUT / output_name).write_text(render_document(SOURCE / source_name, title), encoding='utf-8')
